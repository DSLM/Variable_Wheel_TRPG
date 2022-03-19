# -*- coding: utf-8 -*-
import os
import json
import re
import copy
import string
from docx import Document #pip install python_docx
import lzstring #pip install lzstring
import openpyxl #pip install openpyxl
import xlwings #pip install xlwings
from PIL import Image
import io

#初始化
levelToNum = {"黑色":1,"绿色":2,"蓝色":3,"紫色":4,"红色":5,"橙色":6,"金色":7}
improObj = {"subkeys":{}, "lessArr":{"files":[], "skills":[]}, "moreArr":{"files":[], "skills":[]}}
itemObj = {"subkeys":{}, "lessArr":{"items":[]}, "moreArr":{"items":[]}}
totalSkills = 0
improTrueKeysList = {"抵点":copy.deepcopy(improObj), "效果":copy.deepcopy(improObj), "未分类":copy.deepcopy(improObj)}
totalItems = 0
itemTrueKeysList = {"抵点":copy.deepcopy(itemObj), "效果":copy.deepcopy(itemObj), "未分类":copy.deepcopy(itemObj)}
#更改：生命恢复，能量恢复，远击增幅
offsetKeys = [
"技能冷却时间","技能消耗能量","技能引导消耗","心无恶垢","不杀之刃","永不退避","永不背弃","机能缺损","临时虚弱","反噬","不稳定","誓约武具","严苛环境","乾坤一掷","连段携用","持续占用","条件响应","背水一战","乘胜追击","濒死反击","压倒","独行者","奉献","适应性战术"]
#武具强化是技能词条
effectKeys = [
"技能生效范围","技能生效目标数量","减益效果持续时间","增益效果持续时间","技能释放动作","攻击技能","增幅","技能增幅","稳固攻势","致命","伤害增幅","甲胄反击","封锁","生命栓锁","护甲击破","护甲削减","技能打断","目标锁定","痛击","惊骇","晕击","麻痹打击","寒击","放血","灼烧","注毒","能量抽取","能量燃烧","缓速","强制位移","迟钝","禁令","命中增幅","稳定格挡","坚壁防御","强韧之甲","护甲","连锁护甲","适应性护甲","专项防护","千层甲","甲胄融合","庇护装甲","闪避/防御增幅","守护屏障","防护措施","第六感","心智壁垒","强健体魄","抗性提升","战斗续航","极限顽强","困兽之斗","生存欲望","无惧伤痛","愈合","生命恢复","瞬间恢复","治疗增幅","复苏之风","生命窃取","过量治疗","能量恢复","快速冥想","快速恢复","净化","驱散","属性增幅","能量池增幅","远击","远击增幅","迅捷","熟练","专家","皮糙肉厚","反射强化","疾行","水下呼吸","跃升","负重增幅","位移","学识渊博","收容","技能容纳","透体","光明化","黑暗化","安魂曲","灵媒","灵击","飞行","隐身","意志解放","愈战愈勇","随机应变","蓄势待发","幸运女神的眷顾","技能强化","高速神言","逆境行者","命运重铸","显赫武具","从属物","组合","武具强化","缄默死神","载具用武器","载具用护盾","载具用插件","外骨骼机甲","充能","智慧核心","飞行载具","共同驾驶","变形出击","全能机体","载具用武器","载具用护盾","载具用插件","过载运行","形态限定"]
for key in offsetKeys:
    improTrueKeysList["抵点"]["subkeys"][key] = copy.deepcopy(improObj)
    itemTrueKeysList["抵点"]["subkeys"][key] = copy.deepcopy(itemObj)
for key in effectKeys:
    improTrueKeysList["效果"]["subkeys"][key] = copy.deepcopy(improObj)
    itemTrueKeysList["效果"]["subkeys"][key] = copy.deepcopy(itemObj)

#去除后缀
dataSource = "./DataSource"
improDataBase = "变量之轮强化序列库"
itemDataBase = "变量之轮道具序列库"
for file in os.listdir(dataSource):
    if ("变量之轮词条目录" in file) and (".docx" in file):
        if (file not in "变量之轮词条目录.docx"):
            os.rename(os.path.join(dataSource, file), os.path.join(dataSource, "变量之轮词条目录.docx"))
    if (improDataBase in file) and (file not in improDataBase):
        os.rename(os.path.join(dataSource, file), os.path.join(dataSource, improDataBase))
    if (itemDataBase in file) and (file not in itemDataBase):
        os.rename(os.path.join(dataSource, file), os.path.join(dataSource, itemDataBase))
improDataBase = dataSource + "/" + improDataBase
itemDataBase = dataSource + "/" + itemDataBase


#词条预提取
doc = Document(dataSource + "/变量之轮词条目录.docx")
theFile = open(dataSource + "/变量之轮词条目录.json", 'w', encoding="utf-8")
for paragraph in doc.paragraphs:
    docLines = re.findall(r"^\[([^\]]*)\](.*)", paragraph.text)
    for i in docLines:
        theFile.write(f"'{i[0]}',")
theFile.write(f"\n")
for paragraph in doc.paragraphs:
    docLines = re.findall(r"^\[([^\]]*)\](.*)", paragraph.text)
    for i in docLines:
        theFile.write(f"'{i[0]}':'[{i[0]}]{i[1]}',\n")
theFile.close()

#添加关键词
def pushImproKeys(typeId, fileId, skillId, key, data, subObj):
    #是否有子词条
    if "-" in key:
        keys = key.split('-', 1)
        if keys[0] not in data["subkeys"]:
            data["subkeys"][keys[0]] = copy.deepcopy(subObj)
        #包含子类
        if fileId not in data["subkeys"][keys[0]]["moreArr"]["files"]:
            data["subkeys"][keys[0]]["moreArr"]["files"].append(fileId)
        if skillId not in data["subkeys"][keys[0]]["moreArr"]["skills"]:
            data["subkeys"][keys[0]]["moreArr"]["skills"].append(skillId)
        #处理子词条
        data["subkeys"][keys[0]] = pushImproKeys(typeId, fileId, skillId, keys[1], data["subkeys"][keys[0]], subObj)
    else:
        if key not in data["subkeys"]:
            data["subkeys"][key] = copy.deepcopy(subObj)
        #不包含子类
        if fileId not in data["subkeys"][key]["lessArr"]["files"]:
            data["subkeys"][key]["lessArr"]["files"].append(fileId)
        if skillId not in data["subkeys"][key]["lessArr"]["skills"]:
            data["subkeys"][key]["lessArr"]["skills"].append(skillId)
        #包含子类
        if fileId not in data["subkeys"][key]["moreArr"]["files"]:
            data["subkeys"][key]["moreArr"]["files"].append(fileId)
        if skillId not in data["subkeys"][key]["moreArr"]["skills"]:
            data["subkeys"][key]["moreArr"]["skills"].append(skillId)
    return data

#重构关键词
def rebuildImproKeys(obj):
    for key in obj:
        obj[key]["subkeys"] = rebuildImproKeys(obj[key]["subkeys"])
        if len(obj[key]["subkeys"]) > 0:
            obj[key]["subkeys"].append({"key":"未分类", "subkeys":[], "moreArr":{"files":obj[key]["lessArr"]["files"], "skills":obj[key]["lessArr"]["skills"]}})
            obj[key]["subkeys"].append({"key":"所有", "subkeys":[], "moreArr":{"files":obj[key]["moreArr"]["files"], "skills":obj[key]["moreArr"]["skills"]}})

    tempArray = []
    for key in obj:
        temp = obj[key]
        temp["key"] = key
        tempArray.append(temp)

    return tempArray

#对单文件处理
def dealImproFile(data, typeName, typeId, fileName, fileData, fileId, obj):
    splitData = re.split(r"\n-----*\n", fileData)
    skillRegex = r"(技能名称：)(.*)"
    prioRegex = r"(优先等级：)([0-9]*)"
    realRegex = r"(实际耗点：)(.*)"
    keyRegex = r"([0-9]*)（(抵点：)?([^）]*)）"

    obj["id"] = fileId
    obj["len"] = 0
    obj["data"] = ""

    if typeName == "种族":
        for i in splitData:
            if i == "":
                continue
            obj["data"] += f"<div>{i}</div>"
        return

    elif typeName == "初始技能与日常技能":

        if fileName == "初始技能表":
            skillRegex = r"(\[[^\]]*\])："
        elif fileName == "日常技能（联动表）":
            skillRegex = r"(【[^\]]*】)"

        for i in splitData:
            if i == "":
                continue

            #提取名称
            name = re.findall(skillRegex, i)

            #高亮
            if  len(name) > 0:

                #统计技能总数
                obj["len"] += 1

                #高亮名字
                if fileName == "初始技能表" or fileName == "日常技能（联动表）":
                    repStr = i.replace(name[0], f"""<span class="skillTitle">{name[0]}</span>""")
                else:
                    repStr = i.replace(name[0][0]+name[0][1], f"""<span class="preIntro">{name[0][0]}</span><span class="skillTitle">{name[0][1]}</span>""")
                tempStr = f"""<div id="skill_{typeId}_{fileId}_{obj["len"]}">{repStr}</div>"""
            else:
                tempStr = f"<div>{i}</div>"
            obj["data"] += tempStr
        return

    #技能表，特质，职业
    else:
        for i in splitData:
            if i == "":
                continue

            #提取名称
            name = re.findall(skillRegex, i)

            #高亮
            if  len(name) > 0:

                #统计技能总数
                global totalSkills
                totalSkills += 1
                obj["len"] += 1
                fileEleId = f"""file_{typeId}_{fileId}"""
                skillEleId = f"""skill_{typeId}_{fileId}_{obj["len"]}"""

                #多个名称，有问题
                if  len(name) > 1:
                    print("多个名称，有问题：", fileName, name)

                #高亮名字
                repStr = i.replace(name[0][0]+name[0][1], f"""<span class="preIntro">{name[0][0]}</span><span class="skillTitle">{name[0][1]}</span>""")

                #高亮优先等级
                prioObj = re.findall(prioRegex, repStr)
                if  len(prioObj) > 0:
                    repStr = repStr.replace(prioObj[0][0]+prioObj[0][1], f"""<span class="preIntro">{prioObj[0][0]}</span><span class="prioNum">{prioObj[0][1]}</span>""")

                #提取并替换真实关键词
                real = re.findall(realRegex, i)
                if  len(real) > 0:
                    #多个实际耗点行，有问题
                    if len(real) > 1:
                        print("多个实际耗点行，有问题：", fileName, name)

                    #提取并替换真实关键词
                    newString = f"""<span class="preIntro">{real[0][0]}</span>"""
                    nowIndex = 0
                    for match_obj in re.finditer(keyRegex, real[0][1]):
                        #替换
                        cost = match_obj.group(1)
                        offs = match_obj.group(2)
                        key = match_obj.group(3)
                        newString += real[0][1][nowIndex:match_obj.start()]
                        nowIndex = match_obj.end()
                        #提取
                        #关键词->技能
                        if offs is not None:
                            data["抵点"] = pushImproKeys(typeId, fileEleId, skillEleId, key, data["抵点"], improObj)
                            if fileEleId not in data["抵点"]["moreArr"]["files"]:
                                data["抵点"]["moreArr"]["files"].append(fileEleId)
                            if skillEleId not in data["抵点"]["moreArr"]["skills"]:
                                data["抵点"]["moreArr"]["skills"].append(skillEleId)
                        else:
                            data["效果"] = pushImproKeys(typeId, fileEleId, skillEleId, key, data["效果"], improObj)
                            if fileEleId not in data["效果"]["moreArr"]["files"]:
                                data["效果"]["moreArr"]["files"].append(fileEleId)
                            if skillEleId not in data["效果"]["moreArr"]["skills"]:
                                data["效果"]["moreArr"]["skills"].append(skillEleId)
                        #高亮实际耗点
                        if offs is None:
                            offs = ""
                        else:
                            offs = f"""<span class="preOffset">{offs}</span>"""
                        newString += f"""<span class="prioNum">{cost}</span><span class='fullKey' data-cost='{cost}' data-key='{key}'>（{offs}{key}）</span>"""
                    newString += real[0][1][nowIndex:]

                    #补充遗漏数字高亮
                    restNum = re.findall(r"([0-9]{1,5})(=)", newString)
                    if  len(restNum) > 0:
                        newString = newString.replace(restNum[0][0]+restNum[0][1], f"""<span class="prioNum">{restNum[0][0]}</span>{restNum[0][1]}""")
                    restNum = re.findall(r"(=)([0-9]{1,5})", newString)
                    if  len(restNum) > 0:
                        newString = newString.replace(restNum[0][0]+restNum[0][1], f"""{restNum[0][0]}<span class="prioNum">{restNum[0][1]}</span>""")

                    repStr = repStr.replace(real[0][0]+real[0][1], newString)

                    keys = re.findall(r"([0-9]*)（([^）]*)）", real[0][1])
                    if len(keys) <= 0:
                        #元素id
                        improTrueKeysList["未分类"]["moreArr"]["files"].append(fileEleId)
                        improTrueKeysList["未分类"]["moreArr"]["skills"].append(skillEleId)
                        # DEBUG: 所有无实际耗点的技能
                        print("无实际耗点，有问题：", fileName, name)
                else:
                    #元素id
                    improTrueKeysList["未分类"]["moreArr"]["files"].append(fileEleId)
                    improTrueKeysList["未分类"]["moreArr"]["skills"].append(skillEleId)
                    # DEBUG: 所有无实际耗点的技能
                    print("无实际耗点，有问题：", fileName, name)

                #最后套div
                tempStr = f"""<div id="skill_{typeId}_{fileId}_{obj["len"]}" class="skillText">{repStr}</div>"""
            else:
                tempStr = f"<div>{i}</div>"
            obj["data"] += tempStr

#强化库
improList = {
"初始技能与日常技能":{"id":1},
"技能表":{"id":2},
"特质":{"id":3},
"职业":{"id":4},
"种族":{"id":5}}

#文件导入
for file in os.listdir(improDataBase):
    if ("更新日志" in file):
        if (file not in "更新日志.txt"):
            os.rename(os.path.join(improDataBase, file), os.path.join(improDataBase, "更新日志.txt"))
    for key in improList:
        if (key in file) and (file not in key):
            os.rename(os.path.join(improDataBase, file), os.path.join(improDataBase, key))
for key in improList:
    improList[key]["path"] = improDataBase + "/" + key

#挨个处理文件
for key in improList:
    count = 1
    improList[key]["files"] = {}
    for file in os.listdir(improList[key]["path"]):
        trueName = re.sub(r"\[[^\]]*\]", "", file.replace(".txt", ""))
        improList[key]["files"][trueName] = {}
        f = open(improList[key]["path"] + "/" + file, 'r', encoding="utf-8")
        dealImproFile(improTrueKeysList, key, improList[key]["id"], trueName, f.read(), count, improList[key]["files"][trueName])
        count += 1
    del improList[key]["path"]

#文本分析
newImproList = {"name":"强化序列库", "data":[]}
for type in improList:
        tempFolder = {"name":type, "data":[]}
        for file in improList[type]["files"]:
                key = f"""file_{improList[type]["id"]}_{improList[type]["files"][file]["id"]}"""
                tempFile = {"name":file, "key":key, "len":improList[type]["files"][file]["len"], "data":f"""<div class="title2">{file}</div>{improList[type]["files"][file]["data"]}"""};
                tempFolder["data"].append(tempFile)
        newImproList["data"].append(tempFolder)

#重构关键词
rebuildImproKeysList = rebuildImproKeys(improTrueKeysList)

theFile = open('./src/data/improList.js', 'w', encoding="utf-8")
x = lzstring.LZString()
improListStr = json.dumps(newImproList, ensure_ascii=False)
improTrueKeysListStr = json.dumps(rebuildImproKeysList, ensure_ascii=False)
theFile.write(f"""export var skillsNum = {totalSkills};export var decImproTrueKeysList = '{x.compressToBase64(improTrueKeysListStr)}';export var decImproList = '{x.compressToBase64(improListStr)}';""")
theFile.close()


#添加关键词
def pushItemKeys(id, key, data, subObj):
    #是否有子词条
    if "-" in key:
        keys = key.split('-', 1)
        if keys[0] not in data["subkeys"]:
            data["subkeys"][keys[0]] = copy.deepcopy(subObj)
        #包含子类
        if id not in data["subkeys"][keys[0]]["moreArr"]["items"]:
            data["subkeys"][keys[0]]["moreArr"]["items"].append(id)
        #处理子词条
        data["subkeys"][keys[0]] = pushItemKeys(id, keys[1], data["subkeys"][keys[0]], subObj)
    else:
        if key not in data["subkeys"]:
            data["subkeys"][key] = copy.deepcopy(subObj)
        #不包含子类
        if id not in data["subkeys"][key]["lessArr"]["items"]:
            data["subkeys"][key]["lessArr"]["items"].append(id)
        #包含子类
        if id not in data["subkeys"][key]["moreArr"]["items"]:
            data["subkeys"][key]["moreArr"]["items"].append(id)
    return data

#重构关键词
def rebuildItemKeys(obj):
    for key in obj:
        obj[key]["subkeys"] = rebuildItemKeys(obj[key]["subkeys"])
        if len(obj[key]["subkeys"]) > 0:
            obj[key]["subkeys"].append({"key":"未分类", "subkeys":[], "moreArr":{"items":obj[key]["lessArr"]["items"]}})
            obj[key]["subkeys"].append({"key":"所有", "subkeys":[], "moreArr":{"items":obj[key]["moreArr"]["items"]}})

    tempArray = []
    for key in obj:
        temp = obj[key]
        temp["key"] = key
        tempArray.append(temp)

    return tempArray

#效果文本处理
def dealItemEffect(id, str):
    global itemTrueKeysList
    twoParts = re.split(r"【实际耗点】", str, 1)
    if len(twoParts) != 2:
        #元素id
        itemTrueKeysList["未分类"]["moreArr"]["items"].append(id)
        # DEBUG: 所有无实际耗点的技能
        #print("道具无有耗点？：", id, str)
        return str
    realRegex = r"([^：]*：)(.*)"
    keyRegex = r"([0-9]*)（(抵点：)?([^）]*)）"

    tempRealStr = twoParts[1]

    #提取多个实际耗点
    lines = re.findall(realRegex, twoParts[1])
    if len(lines) == 0:
        #元素id
        itemTrueKeysList["未分类"]["moreArr"]["items"].append(id)
        # DEBUG: 所有无实际耗点的技能
        #print("道具无有耗点？：", id, str)
    for line in lines:
        #提取并替换真实关键词
        newString = f"""<span class="skillTitle">{line[0]}</span>"""
        nowIndex = 0
        for match_obj in re.finditer(keyRegex, line[1]):
            #替换
            cost = match_obj.group(1)
            offs = match_obj.group(2)
            key = match_obj.group(3)
            newString += line[1][nowIndex:match_obj.start()]
            nowIndex = match_obj.end()
            #提取
            #关键词->技能
            if offs is not None:
                itemTrueKeysList["抵点"] = pushItemKeys(id, key, itemTrueKeysList["抵点"], itemObj)
                if id not in itemTrueKeysList["抵点"]["moreArr"]["items"]:
                    itemTrueKeysList["抵点"]["moreArr"]["items"].append(id)
            else:
                itemTrueKeysList["效果"] = pushItemKeys(id, key, itemTrueKeysList["效果"], itemObj)
                if id not in itemTrueKeysList["效果"]["moreArr"]["items"]:
                    itemTrueKeysList["效果"]["moreArr"]["items"].append(id)
            #高亮实际耗点
            if offs is None:
                offs = ""
            else:
                offs = f"""<span class="preOffset">{offs}</span>"""
            newString += f"""<span class="prioNum">{cost}</span><span class='fullKey' data-cost='{cost}' data-key='{key}'>（{offs}{key}）</span>"""
        newString += line[1][nowIndex:]

        #补充遗漏数字高亮
        restNum = re.findall(r"([0-9]{1,5})(=)", newString)
        if  len(restNum) > 0:
            newString = newString.replace(restNum[0][0]+restNum[0][1], f"""<span class="prioNum">{restNum[0][0]}</span>{restNum[0][1]}""")
        restNum = re.findall(r"(=)([0-9]{1,5})", newString)
        if  len(restNum) > 0:
            newString = newString.replace(restNum[0][0]+restNum[0][1], f"""{restNum[0][0]}<span class="prioNum">{restNum[0][1]}</span>""")

        tempRealStr = tempRealStr.replace(line[0] + line[1], newString)

    #返回值
    return twoParts[0] + f"""<span class="preIntro">【实际耗点】</span>""" + tempRealStr


#道具分门别类处理
#武器类型，资历值加成：无用
#近战，魔导
def itemMelee(ws, row, col, id):
    tempItem = {"key": id, "type": "melee"}
    tempItem["名称"] = ws.range(row + 0, col + 1).value
    tempItem["武器大类"] = [
        ws.range(row + 1, col + 1).value,
        ws.range(row + 1, col + 2).value,
        ws.range(row + 1, col + 3).value
    ]
    tempItem["攻击力结算"] = ws.range(row + 2, col + 1).value
    tempItem["攻击力耗点"] = ws.range(row + 2, col + 3).value
    tempItem["品质"] = ws.range(row + 3, col + 1).value
    tempItem["level"] = levelToNum[tempItem["品质"]]
    tempItem["重量"] = ws.range(row + 3, col + 3).value
    tempItem["总计耗点"] = ws.range(row + 4, col + 1).value
    tempItem["体积"] = ws.range(row + 4, col + 3).value
    tempItem["价格"] = ws.range(row + 5, col + 1).value
    tempItem["效果"] = dealItemEffect(id, ws.range(row + 6, col + 0).value)
    tempItem["简介"] = ws.range(row + 6, col + 2).value
    tempItem["效果耗点"] = ws.range(row + 12, col + 1).value
    tempItem["立绘"] = ws.range(row + 13, col + 0).value
    tempItem["制作人"] = ws.range(row + 23, col + 0).value
    tempItem["cost"] = tempItem["总计耗点"]
    return tempItem

#远程
def itemRange(ws, row, col, id):
    tempItem = {"key": id, "type": "range"}
    tempItem["名称"] = ws.range(row + 0, col + 1).value
    tempItem["武器大类"] = [
        ws.range(row + 1, col + 1).value,
        ws.range(row + 1, col + 2).value,
        ws.range(row + 1, col + 3).value
    ]
    tempItem["攻击力结算"] = ws.range(row + 2, col + 1).value
    tempItem["攻击力耗点"] = ws.range(row + 2, col + 3).value
    tempItem["载弹量"] = ws.range(row + 3, col + 1).value
    tempItem["载弹耗点"] = ws.range(row + 3, col + 3).value
    tempItem["射程"] = ws.range(row + 4, col + 1).value
    tempItem["射程耗点"] = ws.range(row + 4, col + 3).value
    tempItem["品质"] = ws.range(row + 5, col + 1).value
    tempItem["level"] = levelToNum[tempItem["品质"]]
    tempItem["重量"] = ws.range(row + 5, col + 3).value
    tempItem["总计耗点"] = ws.range(row + 6, col + 1).value
    tempItem["体积"] = ws.range(row + 6, col + 3).value
    tempItem["价格"] = ws.range(row + 7, col + 1).value
    tempItem["效果"] = dealItemEffect(id, ws.range(row + 8, col + 0).value)
    tempItem["简介"] = ws.range(row + 8, col + 2).value
    tempItem["效果耗点"] = ws.range(row + 14, col + 1).value
    tempItem["立绘"] = ws.range(row + 15, col + 0).value
    tempItem["制作人"] = ws.range(row + 23, col + 0).value
    tempItem["cost"] = tempItem["总计耗点"]
    return tempItem

#防具
def itemArmor(ws, row, col, id):
    tempItem = {"key": id, "type": "armor"}
    tempItem["名称"] = ws.range(row + 0, col + 1).value
    tempItem["防具部位"] = ws.range(row + 0, col + 3).value
    tempItem["防御力等级"] = ws.range(row + 1, col + 1).value
    tempItem["防御力耗点"] = ws.range(row + 1, col + 3).value
    tempItem["品质"] = ws.range(row + 2, col + 1).value
    tempItem["level"] = levelToNum[tempItem["品质"]]
    tempItem["重量"] = ws.range(row + 2, col + 3).value
    tempItem["总计耗点"] = ws.range(row + 3, col + 1).value
    tempItem["体积"] = ws.range(row + 3, col + 3).value
    tempItem["价格"] = ws.range(row + 4, col + 1).value
    tempItem["效果"] = dealItemEffect(id, ws.range(row + 5, col + 0).value)
    tempItem["简介"] = ws.range(row + 5, col + 2).value
    tempItem["效果耗点"] = ws.range(row + 11, col + 1).value
    tempItem["立绘"] = ws.range(row + 12, col + 0).value
    tempItem["制作人"] = ws.range(row + 23, col + 0).value
    tempItem["cost"] = tempItem["总计耗点"]
    return tempItem

#盾牌
def itemShield(ws, row, col, id):
    tempItem = {"key": id, "type": "shield"}
    tempItem["名称"] = ws.range(row + 0, col + 1).value
    tempItem["具体分类"] = ws.range(row + 0, col + 3).value
    tempItem["防御力结算"] = ws.range(row + 1, col + 1).value
    tempItem["防御力耗点"] = ws.range(row + 1, col + 3).value
    tempItem["品质"] = ws.range(row + 2, col + 1).value
    tempItem["level"] = levelToNum[tempItem["品质"]]
    tempItem["重量"] = ws.range(row + 2, col + 3).value
    tempItem["总计耗点"] = ws.range(row + 3, col + 1).value
    tempItem["体积"] = ws.range(row + 3, col + 3).value
    tempItem["价格"] = ws.range(row + 4, col + 1).value
    tempItem["效果"] = dealItemEffect(id, ws.range(row + 5, col + 0).value)
    tempItem["简介"] = ws.range(row + 5, col + 2).value
    tempItem["效果耗点"] = ws.range(row + 11, col + 1).value
    tempItem["立绘"] = ws.range(row + 12, col + 0).value
    tempItem["制作人"] = ws.range(row + 23, col + 0).value
    tempItem["cost"] = tempItem["总计耗点"]
    return tempItem

#工具
def itemTool(ws, row, col, id):
    tempItem = {"key": id, "type": "tool"}
    tempItem["名称"] = ws.range(row + 0, col + 1).value
    tempItem["品质"] = ws.range(row + 0, col + 3).value
    tempItem["level"] = levelToNum[tempItem["品质"]]
    tempItem["效果"] = dealItemEffect(id, ws.range(row + 1, col + 0).value)
    tempItem["重量"] = ws.range(row + 1, col + 3).value
    tempItem["体积"] = ws.range(row + 2, col + 3).value
    tempItem["简介"] = ws.range(row + 3, col + 2).value
    tempItem["效果耗点"] = ws.range(row + 8, col + 1).value
    tempItem["价格"] = ws.range(row + 9, col + 1).value
    tempItem["立绘"] = ws.range(row + 10, col + 0).value
    tempItem["制作人"] = ws.range(row + 23, col + 0).value
    tempItem["cost"] = tempItem["效果耗点"]
    return tempItem

#载具
def itemCar(ws, row, col, id):
    tempItem = {"key": id, "type": "car"}
    tempItem["名称"] = ws.range(row + 0, col + 1).value
    tempItem["载具类型"] = ws.range(row + 0, col + 5).value

    tempItem["基础防护"] = ws.range(row + 1, col + 1).value
    tempItem["结构强度"] = ws.range(row + 2, col + 1).value
    tempItem["载具出力"] = ws.range(row + 3, col + 1).value
    tempItem["能量值"] = ws.range(row + 4, col + 1).value

    tempItem["防护耗点"] = ws.range(row + 1, col + 3).value
    tempItem["结构耗点"] = ws.range(row + 2, col + 3).value
    tempItem["出力耗点"] = ws.range(row + 3, col + 3).value
    tempItem["能量耗点"] = ws.range(row + 4, col + 3).value

    tempItem["机动补正"] = ws.range(row + 5, col + 1).value
    tempItem["耐久补正"] = ws.range(row + 6, col + 1).value
    tempItem["耗能补正"] = ws.range(row + 7, col + 1).value
    tempItem["保护补正"] = ws.range(row + 5, col + 3).value
    tempItem["灵活补正"] = ws.range(row + 6, col + 3).value
    tempItem["呼叫补正"] = ws.range(row + 7, col + 3).value

    tempItem["载具防护"] = ws.range(row + 8, col + 1).value
    tempItem["载具速度"] = ws.range(row + 9, col + 1).value
    tempItem["载具耗能"] = ws.range(row + 10, col + 1).value
    tempItem["保护效果"] = ws.range(row + 8, col + 3).value
    tempItem["载具时速"] = ws.range(row + 9, col + 3).value
    tempItem["载具耐久"] = ws.range(row + 10, col + 3).value

    tempItem["恶劣环境"] = ws.range(row + 2, col + 5).value
    tempItem["加速性能"] = ws.range(row + 3, col + 5).value
    tempItem["动力类型"] = ws.range(row + 4, col + 5).value
    tempItem["乘员上限"] = ws.range(row + 5, col + 5).value
    tempItem["灵活性"] = ws.range(row + 2, col + 7).value
    tempItem["乘员防护"] = ws.range(row + 3, col + 7).value
    tempItem["呼叫时间"] = ws.range(row + 4, col + 7).value
    tempItem["负重上限"] = ws.range(row + 5, col + 7).value

    tempItem["品质"] = ws.range(row + 11, col + 1).value
    tempItem["level"] = levelToNum[tempItem["品质"]]
    tempItem["重量"] = ws.range(row + 11, col + 3).value
    tempItem["总计耗点"] = ws.range(row + 12, col + 1).value
    tempItem["体积"] = ws.range(row + 12, col + 3).value
    tempItem["价格"] = ws.range(row + 13, col + 1).value

    tempItem["效果"] = dealItemEffect(id, ws.range(row + 14, col + 0).value)
    tempItem["效果耗点"] = ws.range(row + 21, col + 2).value
    tempItem["简介"] = ws.range(row + 6, col + 4).value
    tempItem["立绘"] = ws.range(row + 14, col + 4).value

    tempItem["制作人"] = ws.range(row + 22, col + 0).value
    tempItem["cost"] = tempItem["总计耗点"]
    return tempItem

#图片导出
def excelImage(images, row_start, col_start, row_end, col_end, id):
    #记得删，无需二次
    return
    for row in range(row_start, row_end + 1):
        for col in range(col_start, col_end + 1):
            strCoor = f'{col}_{row}'
            if strCoor in images:
                itemsImage.append(id)
                images[strCoor].save(rf'./src/data/itemsImage/{id}.png')
                return

#排序方法
def itemSortGetCost(ele):
    return int(ele["cost"])
def itemSortByCost(x, y):
    if int(x["cost"]) > int(y["cost"]):
        return 1
    elif int(x["cost"]) < int(y["cost"]):
        return -1
    else:
        return 0

#道具库
itemsImage = []
itemList = {
"变量军械库":{"id":1},
"变量防具库":{"id":2},
"变量工具库":{"id":3}}
itemMenusList = {
"变量军械库":["剑", "刀", "拳套", "长柄", "斧锤", "奇门兵器", "弓", "弩", "半自动枪械", "全自动枪械", "非自动枪械", "魔导器", "共生体武器", "副武器"],
"变量防具库":["头部", "身体", "背部", "手臂", "腰部", "腿部", "饰品", "共生体", "盾牌", "背包"],
"变量工具库":["恢复类", "造伤类", "消耗类", "工具类", "组合配件", "载具"]}

#文件导入
for file in os.listdir(itemDataBase):
    if ("更新日志" in file):
        if (file not in "更新日志.txt"):
            os.rename(os.path.join(itemDataBase, file), os.path.join(itemDataBase, "更新日志.txt"))
    for key in itemList:
        if (key in file) and (file not in key):
            os.rename(os.path.join(itemDataBase, file), os.path.join(itemDataBase, key+".xlsx"))
for file in itemList:
    itemList[file]["path"] = itemDataBase + "/" + file + ".xlsx"

#挨个处理道具
#遍历文件
for file in itemList:
    menuId = 1
    itemList[file]["menus"] = {}
    wb = xlwings.Book(itemList[file]["path"])
    wb_op = openpyxl.load_workbook(itemList[file]["path"])
    print(file)

    #遍历工作薄
    for menu in itemMenusList[file]:
        itemId = 1
        itemList[file]["menus"][menu] = {"name": menu, "id": menuId, "items": []}
        ws = wb.sheets[menu]
        ws_op = wb_op[menu]
        #自行实现图片导入
        images = {}
        sheet_images = ws_op._images
        for image in sheet_images:
            row = image.anchor._from.row + 1
            col = image.anchor._from.col + 1
            images[f'{col}_{row}'] = Image.open(io.BytesIO(image._data()))
        #遍历单元格
        for row in ws_op.iter_rows():
            for cell in row:
                #这是一个有名字的东西
                if cell.value == "名称":
                    #元素id
                    strId = f"""item_{itemList[file]["id"]}_{menuId}_{itemId}"""
                    if ws_op.cell(cell.row + 4, cell.column + 0).value == "总计耗点":
                        #近战，魔导
                        itemList[file]["menus"][menu]["items"].append(itemMelee(ws, cell.row, cell.column, strId))
                        excelImage(images, cell.row, cell.column - 1, cell.row + 23, cell.column + 3, strId)
                    elif ws_op.cell(cell.row + 4, cell.column + 0).value == "射程":
                        #远程
                        itemList[file]["menus"][menu]["items"].append(itemRange(ws, cell.row, cell.column, strId))
                        excelImage(images, cell.row, cell.column - 1, cell.row + 23, cell.column + 3, strId)
                    elif ws_op.cell(cell.row + 1, cell.column + 0).value == "防御力等级":
                        #防具
                        itemList[file]["menus"][menu]["items"].append(itemArmor(ws, cell.row, cell.column, strId))
                        excelImage(images, cell.row, cell.column - 1, cell.row + 23, cell.column + 3, strId)
                    elif ws_op.cell(cell.row + 1, cell.column + 0).value == "防御力结算":
                        #盾牌
                        itemList[file]["menus"][menu]["items"].append(itemShield(ws, cell.row, cell.column, strId))
                        excelImage(images, cell.row, cell.column - 1, cell.row + 23, cell.column + 3, strId)
                    elif ws_op.cell(cell.row + 0, cell.column + 2).value == "品质":
                        #工具
                        itemList[file]["menus"][menu]["items"].append(itemTool(ws, cell.row, cell.column, strId))
                        excelImage(images, cell.row, cell.column - 1, cell.row + 23, cell.column + 3, strId)
                    elif ws_op.cell(cell.row + 0, cell.column + 4).value == "载具类型":
                        #载具
                        itemList[file]["menus"][menu]["items"].append(itemCar(ws, cell.row, cell.column, strId))
                        excelImage(images, cell.row, cell.column - 1, cell.row + 22, cell.column + 8, strId)
                    else:
                        #有问题
                        print(ws_op.cell(cell.row + 0, cell.column + 1).value)
                        itemId -= 1
                        totalItems -=1
                    itemId += 1
                    totalItems += 1
        menuId += 1
    del itemList[file]["path"]
    wb.save()
    wb.close()

#文本分析
newItemList = []
for file in itemList:
    for menu in itemMenusList[file]:
        itemList[file]["menus"][menu]["items"].sort(key=itemSortGetCost)
    tempFile = {"name":file, "menus":itemList[file]["menus"]}
    newItemList.append(tempFile)

#重构关键词
rebuildItemKeysList = rebuildItemKeys(itemTrueKeysList)

theFile = open('./src/data/itemList.js', 'w', encoding="utf-8")
x = lzstring.LZString()
itemsImageStr = json.dumps(itemsImage, ensure_ascii=False)
itemListStr = json.dumps(newItemList, ensure_ascii=False)
itemTrueKeysListStr = json.dumps(rebuildItemKeysList, ensure_ascii=False)
theFile.write(f"""export var itemsNum = {totalItems};export var itemsImage = {itemsImageStr};export var decItemTrueKeysList = '{x.compressToBase64(itemTrueKeysListStr)}';export var decItemList = '{x.compressToBase64(itemListStr)}';""")
theFile.close()
print("End")
